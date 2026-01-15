"""Document text extraction for PDF, DOCX, and TXT files."""

from enum import Enum
from io import BytesIO
from pathlib import Path
from typing import BinaryIO

from pydantic import BaseModel, Field


class SupportedDocumentType(str, Enum):
    """Supported document types for extraction."""

    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"


class ExtractionResult(BaseModel):
    """Result of document text extraction."""

    text: str = Field(..., description="Extracted text content")
    char_count: int = Field(..., description="Number of characters extracted")
    page_count: int | None = Field(default=None, description="Number of pages (for PDFs)")
    was_truncated: bool = Field(default=False, description="Whether text was truncated")
    truncation_reason: str | None = Field(default=None, description="Reason for truncation")
    warnings: list[str] = Field(default_factory=list, description="Extraction warnings")


class DocumentExtractor:
    """Extract text from various document formats.

    Supports PDF, DOCX, and TXT files. Runs locally without external API calls.
    """

    MAX_CHARS = 10000  # Match ProjectInput.description max_length
    MAX_PDF_PAGES = 50  # Limit PDF extraction to first 50 pages

    def extract(self, file: BinaryIO, filename: str) -> ExtractionResult:
        """Extract text from an uploaded file.

        Args:
            file: Binary file-like object containing the document.
            filename: Original filename (used to determine file type).

        Returns:
            ExtractionResult with extracted text and metadata.

        Raises:
            ValueError: If file type is unsupported or file cannot be read.
        """
        ext = Path(filename).suffix.lower().lstrip(".")

        if ext == "pdf":
            return self._extract_pdf(file)
        elif ext == "docx":
            return self._extract_docx(file)
        elif ext in ("txt", "md", "text"):
            return self._extract_txt(file)
        else:
            raise ValueError(
                f"Unsupported file type: .{ext}. "
                f"Supported types: PDF, DOCX, TXT"
            )

    def _extract_pdf(self, file: BinaryIO) -> ExtractionResult:
        """Extract text from PDF document."""
        try:
            from pypdf import PdfReader
            from pypdf.errors import PdfReadError
        except ImportError as e:
            raise ValueError(
                "PDF extraction requires pypdf. Install with: pip install pypdf"
            ) from e

        warnings: list[str] = []

        try:
            reader = PdfReader(file)
            total_pages = len(reader.pages)
            pages_to_extract = min(total_pages, self.MAX_PDF_PAGES)

            if total_pages > self.MAX_PDF_PAGES:
                warnings.append(
                    f"PDF has {total_pages} pages. "
                    f"Only first {self.MAX_PDF_PAGES} pages extracted."
                )

            text_parts: list[str] = []
            for i, page in enumerate(reader.pages[:pages_to_extract]):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception:
                    warnings.append(f"Could not extract page {i + 1}")

            if not text_parts:
                raise ValueError(
                    "No text could be extracted from PDF. "
                    "The file may be image-based or encrypted."
                )

            full_text = "\n\n".join(text_parts)
            was_truncated = len(full_text) > self.MAX_CHARS
            final_text = full_text[: self.MAX_CHARS]

            return ExtractionResult(
                text=final_text,
                char_count=len(final_text),
                page_count=pages_to_extract,
                was_truncated=was_truncated,
                truncation_reason=(
                    f"Text exceeded {self.MAX_CHARS:,} character limit"
                    if was_truncated
                    else None
                ),
                warnings=warnings,
            )

        except PdfReadError as e:
            raise ValueError(f"Cannot read PDF: {e}") from e

    def _extract_docx(self, file: BinaryIO) -> ExtractionResult:
        """Extract text from Word document."""
        try:
            from docx import Document
            from docx.opc.exceptions import PackageNotFoundError
        except ImportError as e:
            raise ValueError(
                "Word extraction requires python-docx. "
                "Install with: pip install python-docx"
            ) from e

        warnings: list[str] = []

        try:
            doc = Document(file)

            # Extract text from paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)

            if not paragraphs:
                raise ValueError(
                    "No text could be extracted from Word document. "
                    "The file may be empty or corrupted."
                )

            full_text = "\n\n".join(paragraphs)
            was_truncated = len(full_text) > self.MAX_CHARS
            final_text = full_text[: self.MAX_CHARS]

            return ExtractionResult(
                text=final_text,
                char_count=len(final_text),
                was_truncated=was_truncated,
                truncation_reason=(
                    f"Text exceeded {self.MAX_CHARS:,} character limit"
                    if was_truncated
                    else None
                ),
                warnings=warnings,
            )

        except PackageNotFoundError as e:
            raise ValueError(
                f"Cannot read Word document: {e}. "
                "The file may be corrupted or not a valid DOCX."
            ) from e

    def _extract_txt(self, file: BinaryIO) -> ExtractionResult:
        """Extract text from plain text file."""
        content = file.read()

        # Try common encodings
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252", "iso-8859-1"]
        full_text: str | None = None

        for encoding in encodings:
            try:
                full_text = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue

        if full_text is None:
            raise ValueError(
                "Cannot decode text file. "
                "The file may use an unsupported character encoding."
            )

        # Clean up the text
        full_text = full_text.strip()

        if not full_text:
            raise ValueError("Text file is empty.")

        was_truncated = len(full_text) > self.MAX_CHARS
        final_text = full_text[: self.MAX_CHARS]

        return ExtractionResult(
            text=final_text,
            char_count=len(final_text),
            was_truncated=was_truncated,
            truncation_reason=(
                f"Text exceeded {self.MAX_CHARS:,} character limit"
                if was_truncated
                else None
            ),
        )

    @staticmethod
    def get_supported_extensions() -> list[str]:
        """Get list of supported file extensions."""
        return ["pdf", "docx", "txt", "md"]

    @staticmethod
    def get_supported_mime_types() -> set[str]:
        """Get set of supported MIME types."""
        return {
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "text/plain",
            "text/markdown",
        }
